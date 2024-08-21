import autogen
import pandas as pd
from datetime import datetime
from docx import Document
from io import BytesIO

# Configuration setup
config_list = autogen.config_list_from_json(env_or_file="OAI_CONFIG_LIST")
llm_config = {"config_list": config_list}

# Calculate the ranges of experience in skills, sectors, functions, and tooling
def merge_date_ranges(ranges):
    if not ranges:
        return []

    sorted_ranges = sorted(ranges, key=lambda x: x[0])
    merged_ranges = []
    current_start, current_end = sorted_ranges[0]

    for start, end in sorted_ranges[1:]:
        if start <= current_end:
            current_end = max(current_end, end)
        else:
            merged_ranges.append((current_start, current_end))
            current_start, current_end = start, end
    merged_ranges.append((current_start, current_end))
    
    return merged_ranges

# Calculate the actual years of experience in skills, sectors, functions, and tooling
def calculate_experience(df):
    experience = {
        "skills": {},
        "sector": {},
        "function": {},
        "tools": {}
    }

    def add_experience(experience_dict, category, start_year, end_year):
        if category not in experience_dict:
            experience_dict[category] = []
        experience_dict[category].append((start_year, end_year))

    for index, job in df.iterrows():
        start_year = int(job['Start'])
        end_year = int(job['End']) if job['End'] else datetime.now().year

        # Process skills
        for skill in job['Skills'].split(','):
            skill = skill.strip()
            add_experience(experience['skills'], skill, start_year, end_year)

        # Process sector
        sector = job['Sector'].strip()
        add_experience(experience['sector'], sector, start_year, end_year)

        # Process function
        function = job['Function'].strip()
        add_experience(experience['function'], function, start_year, end_year)

        # Process tools
        for tool in job['Tools_and_technologies_used'].split(','):
            tool = tool.strip()
            add_experience(experience['tools'], tool, start_year, end_year)

    for category in experience:
        for key in experience[category]:
            merged_ranges = merge_date_ranges(experience[category][key])
            unique_years = sum(end - start + 1 for start, end in merged_ranges)
            experience[category][key] = unique_years

    return experience

def generate_motivation_letter(df, job_description):
    
    print("calculate experience")
    # Calculate experience from the DataFrame
    experience = calculate_experience(df)

    # Prepare the resume details
    resume_details = "\n\n".join(
        f"Company: {row['Company']}\n"
        f"Sector: {row['Sector']}\n"
        f"Function: {row['Function']}\n"
        f"Job Description: {row['Job_description']}\n"
        f"Achievements: {row['Achievements']}\n"
        f"Tools and Technologies Used: {row['Tools_and_technologies_used']}\n"
        f"Skills: {row['Skills']}\n"
        f"Start: {row['Start']}\n"
        f"End: {row['End']}"
        for _, row in df.iterrows()
    )

    # Define tasks for agents
    match_task = """
    Your task is to find the matches between the provided job description and the candidate's calculated experience. 
    Please follow these steps:

    1. Analyze the job description to extract key requirements in terms of skills, sectors, functions, and tools.
    2. Compare these requirements with the candidate's experience.
    3. Identify the most relevant matches for each category (skills, sectors, functions, tools).
    4. Provide a detailed list of these matches, including a brief explanation of how each match is relevant to the job description.
    5. Highlight any unique or standout qualifications that make the candidate a strong fit for the position.

    Reply 'MATCH_DONE' at the end of your response.
    """
    
    strategy_task = """
    Develop a comprehensive strategy for writing exceptional motivation and cover letters, considering diverse application contexts such as job applications, university admissions, and grant proposals. The strategy should include:

    1. Methods for identifying the purpose of each letter and aligning it with the organization's objectives.
    2. Guidelines for analyzing the intended audience to tailor the content accurately.
    3. A plan for structuring these letters and personalizing them to highlight the applicant's unique skills, passions, and qualifications. Mention relevant achievements.
    4. Guidance on using persuasive language and maintaining a professional tone.
    5. Approaches for emphasizing the applicant's unique value, focusing on specific skills or experiences that distinguish them from other candidates and benefit the organization.
    6. Best practices for proofreading and editing to ensure error-free, polished documents.
    7. Suggestions for dynamic opening lines and compelling closing statements that make a memorable impact and urge the reader to take action.

    Compile these guidelines into a comprehensive document.

    Reply 'STRATEGY_DONE' at the end of your response.
    """

    motivation_letter_task = """
    Write a compelling motivation letter based on the matches found between the job description and the candidate's experience, 
    and following the provided strategy guidelines. The letter should:

    1. Highlight the most relevant skills, experiences, and achievements of the candidate, with a focus on key achievements and quantifiable results.
    2. Use persuasive and engaging language while maintaining a professional tone.
    3. Be well-structured, starting with a strong, engaging, and memorable opening, detailed body, and a compelling closing.
    4. Reflect the candidate's unique value and how they align with the job requirements and the company's mission and goals.
    5. Include a clear and compelling call to action.
    6. Ensure the content is coherent, clear, and free from grammatical errors.
    7. Tailor the content specifically to the job and company, demonstrating a deep understanding of their needs and culture.
    8. Include relevant keywords from the job description to ensure it passes through Applicant Tracking Systems (ATS).
    9. Showcase the candidate's enthusiasm for the role and how they fit into the company culture.

    Reply 'TERMINATE' at the end of your response.
    """

    quality_check_task = """
    Review the motivation letter for the following criteria:

    1. Clarity: Ensure the content is easy to understand and free from ambiguity.
    2. Coherence: Check that the ideas flow logically and the letter is well-structured.
    3. Grammar: Correct any grammatical, spelling, or punctuation errors.
    4. Adherence to Guidelines: Ensure the letter follows the provided structure and strategy guidelines.
    5. Formatting: Verify the document is well-formatted, with proper sections and bullet points where necessary.
    6. Overall Quality: Assess the letter's overall quality, including the strength of the opening, alignment with the company’s mission and goals, and the effectiveness of the call to action.
    7. Highlighting Achievements: Ensure key achievements and quantifiable results are clearly highlighted.
    8. Tailoring: Verify that the content is specifically tailored to the job and company, demonstrating an understanding of their needs and culture.
    9. Keywords: Check that relevant keywords from the job description are included to ensure it passes through Applicant Tracking Systems (ATS).
    10. Enthusiasm and Cultural Fit: Ensure the candidate's enthusiasm for the role and cultural fit with the company are effectively showcased.

    Provide detailed feedback and necessary corrections. Reply 'QUALITY_DONE' at the end of your review. If revisions are needed, provide specific feedback and reply 'REWRITE_NEEDED' at the very end.
    """

    # Define agents
    match_maker = autogen.AssistantAgent(
        name="Match_Maker",
        llm_config=llm_config,
        system_message="""
            You are an expert match maker for job applications. Your task is to meticulously analyze the provided job description 
            and the candidate's calculated experience, identifying the most relevant skills, sectors, functions, and tools that 
            match the job requirements. Your response should include a detailed list of matches along with a brief explanation 
            of how each match is relevant to the job description. Ensure to highlight any unique or standout qualifications 
            that make the candidate a strong fit for the position. 
            Reply 'MATCH_DONE' at the end of your response.
            """,
    )

    strategy_agent = autogen.AssistantAgent(
        name="Strategy_Agent",
        llm_config=llm_config,
        system_message="""
            You are an expert in crafting effective and persuasive motivation and cover letters tailored for various applications. 
            Your task is to develop an advanced strategy for writing these letters, considering different contexts such as job applications, university admissions, and grant proposals. 
            The strategy should encompass the following areas:
            
            1. **Purpose and Audience Analysis**: Provide methods to identify the purpose of each letter and align it with the organization's goals. 
               Include guidelines for analyzing the intended audience to tailor the content appropriately.
            
            2. **Content Structuring and Personalization**: Suggest a plan for structuring the letters and personalizing them to highlight the applicant's unique skills, passions, and qualifications. 
               Emphasize the importance of mentioning relevant achievements.
            
            3. **Persuasive Language Usage**: Offer guidance on using persuasive language and maintaining a professional tone to effectively engage the reader.
            
            4. **Highlighting Unique Value**: Propose strategies to emphasize the applicant's unique value, focusing on specific skills or experiences that distinguish them from other candidates 
               and offer particular benefits to the organization.
            
            5. **Proofreading and Editing Best Practices**: Recommend best practices for ensuring the documents are error-free and polished.
            
            6. **Effective Openings and Closings**: Provide suggestions for dynamic opening lines and compelling closing statements that make a memorable impact and urge the reader to take action.
            
            Compile these guidelines into a comprehensive document that can serve as a blueprint for drafting motivation and cover letters.
            
            Reply 'STRATEGY_DONE' at the end of your response.
            """,
    )

    writing_agent = autogen.AssistantAgent(
        name="Writing_Agent",
        llm_config=llm_config,
        system_message="""
            You are a professional writer specializing in crafting compelling motivation letters. 
            Your task is to create a motivation letter based on the matches between the job description and the candidate's experience, 
            using the provided strategy guidelines. The letter should be persuasive, well-structured, and tailored to highlight the candidate's 
            unique qualifications and achievements relevant to the job. Ensure the language is engaging and professional, and the content is 
            aligned with the strategy's recommendations.

            Reply 'TERMINATE' at the end of your response.
            """,
    )


    quality_check_agent = autogen.AssistantAgent(
        name="Quality_Check_Agent",
        llm_config=llm_config,
        system_message="""
            You are a quality assurance specialist for professional documents. 
            Your task is to review the motivation letter for clarity, coherence, grammar, and adherence to the provided structure and guidelines. 
            Ensure the letter is well-formatted, with proper sections and bullet points where necessary. 
            Provide detailed feedback and necessary corrections to enhance the letter's overall quality and impact.

            Reply 'QUALITY_DONE' at the end of your review. If revisions are needed, provide specific feedback and reply 'REWRITE_NEEDED' at the very end.
            """,
    )


    # User proxy agent
    user_proxy_auto = autogen.UserProxyAgent(
        name="User_Proxy_Auto",
        human_input_mode="NEVER",
        is_termination_msg=lambda x: x.get("content", "").rstrip().endswith("TERMINATE"),
        code_execution_config={
            "last_n_messages": 1,
            "work_dir": "tasks",
            "use_docker": False,
        },
    )


    # Prepare the experience data
    experience_input = f"""
    Skills: {experience['skills']}
    Sector: {experience['sector']}
    Function: {experience['function']}
    Tools: {experience['tools']}
    """

    # Step 1: Find matches between job description and calculated experience
    match_task_with_experience = f"""
    Job Description: {job_description}
    Calculated Experience: {experience_input}
    """

    match_results = autogen.initiate_chats(
        [
            {
                "sender": user_proxy_auto,
                "recipient": match_maker,
                "message": match_task_with_experience,
                "summary_method": "last_msg",
                "max_turns": 1,
            }
        ]
    )

    match_result = match_results[0].summary

    if "MATCH_DONE" not in match_result:
        raise Exception("Matching process did not complete successfully.")

    # Step 2: Find the best strategy for the motivation letter
    strategy_input = f"""
    Job Description: {job_description}
    Calculated Experience: {experience_input}
    Matches found: {match_result}
    Resume details: {resume_details}
    """

    strategy_results = autogen.initiate_chats(
        [
            {
                "sender": user_proxy_auto,
                "recipient": strategy_agent,
                "message": strategy_task,
                "carryover": strategy_input,
                "summary_method": "last_msg",
                "max_turns": 1,
            }
        ]
    )

    strategy_result = strategy_results[0].summary

    if "STRATEGY_DONE" not in strategy_result:
        raise Exception("Strategy generation did not complete successfully.")

    # Step 3: Write a motivation letter
    motivation_letter_input = f"""
    Matches found: {match_result}
    Strategy suggested: {strategy_result}
    Resume Details: {resume_details}
    """

    motivation_letter_results = autogen.initiate_chats(
        [
            {
                "sender": user_proxy_auto,
                "recipient": writing_agent,
                "message": motivation_letter_task,
                "carryover": motivation_letter_input,
                "max_turns": 1,
            }
        ]
    )

    motivation_letter_result = motivation_letter_results[0].summary

    # Step 4: Quality check the motivation letter
    quality_check_input = f"""
    Motivation Letter: {motivation_letter_result}
    """

    quality_check_results = autogen.initiate_chats(
        [
            {
                "sender": user_proxy_auto,
                "recipient": quality_check_agent,
                "message": quality_check_task,
                "carryover": quality_check_input,
                "summary_method": "last_msg",
                "max_turns": 1,
            }
        ]
    )

    quality_check_result = quality_check_results[0].summary

    if "QUALITY_DONE" not in quality_check_result and "REWRITE_NEEDED" not in quality_check_result:
        raise Exception("Quality check did not complete successfully.")

    if "REWRITE_NEEDED" in quality_check_result:
        feedback = quality_check_result.replace('REWRITE_NEEDED', '').strip()
        
        motivation_letter_results = autogen.initiate_chats(
            [
                {
                    "sender": user_proxy_auto,
                    "recipient": writing_agent,
                    "message": motivation_letter_task,
                    "carryover": motivation_letter_result + f"\n\nQuality Check Feedback: {feedback}",
                    "summary_method": "last_msg",
                    "max_turns": 1,
                }
            ]
        )
        motivation_letter_result = motivation_letter_results[0].summary

    # Save the motivation letter to a .docx file in memory
    document = Document()
    document.add_heading('Motivation Letter', level=1)
    document.add_paragraph(motivation_letter_result)

    # Save the document in-memory rather than to disk
    doc_io = BytesIO()
    document.save(doc_io)
    doc_io.seek(0)

    return doc_io
