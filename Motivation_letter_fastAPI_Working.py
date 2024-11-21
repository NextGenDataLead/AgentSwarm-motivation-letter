from fastapi import FastAPI, File, UploadFile, Form, Depends
from fastapi.responses import RedirectResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import Optional
from io import BytesIO
from docx import Document
import pandas as pd
import os
import datetime
import autogen
from fastapi.middleware.cors import CORSMiddleware
import asyncio
import logging
from sse_starlette.sse import EventSourceResponse
from fastapi.responses import FileResponse

# Setup logging configuration
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s', handlers=[logging.StreamHandler()])
logger = logging.getLogger(__name__)

# Configuration setup
config_list = autogen.config_list_from_json(env_or_file="OAI_CONFIG_LIST")
filtered_config_list = [item for item in config_list if item["model"] == "gpt-4"]
llm_config = {"config_list": filtered_config_list}

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve static files (HTML, CSS, etc.)
app.mount("/static", StaticFiles(directory="static"), name="static")

log_messages = []  # A global list to store logs

# Custom handler for capturing logs in real-time with forced flushing
class LogListHandler(logging.Handler):
    def emit(self, record):
        log_entry = self.format(record)
        log_messages.append(log_entry)
        self.flush()

    def flush(self):
        pass

# Add the custom handler to capture logs
log_handler = LogListHandler()
log_handler.setLevel(logging.INFO)
log_handler.setFormatter(logging.Formatter('%(asctime)s - %(message)s'))
logger.addHandler(log_handler)

# Global dictionary to store session data (for simplicity)
session_data = {
    "user_feedback": None,
    "last_generated_output": None,
    "current_job_description": None
}

# SSE endpoint to stream logs
@app.get("/logs/")
async def stream_logs():
    async def log_generator():
        while True:
            if log_messages:
                message = log_messages.pop(0)
                yield f"data: {message}\n\n"
            await asyncio.sleep(0.1)

    return EventSourceResponse(log_generator())

# Root route that redirects to the static index.html
@app.get("/")
async def serve_index():
    return RedirectResponse(url="/static/index.html")

# Asynchronous resume upload and motivation letter generation
@app.post("/generate-letter/")
async def generate_motivation_letter(
    api_key: str = Form(...), 
    resume_file: UploadFile = File(...), 
    job_description: str = Form(...), 
    user_feedback: Optional[str] = Form(None)  # Accepting feedback as 'user_feedback'
):
    
    global session_data
    
    # Ensure feedback ends with a period
    if user_feedback and not user_feedback.endswith('.'):
        user_feedback += '.'

    # Reset feedback and output if a new job description is provided
    if session_data["current_job_description"] != job_description:
        session_data["current_job_description"] = job_description
        session_data["user_feedback"] = user_feedback
        session_data["last_generated_output"] = None
    else:
        # Append new feedback to existing feedback if additional feedback is provided
        if user_feedback:
            session_data["user_feedback"] = (
                session_data["user_feedback"] + " " + user_feedback if session_data["user_feedback"] else user_feedback
            )

    await asyncio.sleep(1)
    logger.info("Received request to generate motivation letter...")
    await asyncio.sleep(1)

    # Asynchronously read the resume
    resume_df = pd.read_excel(resume_file.file, sheet_name="Jobs")
    await asyncio.sleep(1)
    logger.info("Processing resume file...")
    await asyncio.sleep(1)

    # Run the letter creation process using stored feedback and last output
    motivation_letter_result = await run_letter_creation_process(
        resume_df, 
        job_description, 
        session_data["user_feedback"],
        session_data["last_generated_output"]
    )

    # Update the last generated output in the session data
    session_data["last_generated_output"] = motivation_letter_result

    # Generate motivation letter as a Word document in-memory
    document = Document()
    document.add_heading("Motivation Letter", level=1)
    document.add_paragraph(motivation_letter_result)
    await asyncio.sleep(1)
    logger.info("Generating the motivation letter document...")
    await asyncio.sleep(1)

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    await asyncio.sleep(1)
    logger.info("Motivation letter generation complete!")
    await asyncio.sleep(1)

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=motivation_letter.docx"},
    )

@app.get("/download-template/")
async def download_template():
    # Path to the template file
    template_path = os.path.join("Resume", "TEMPLATE_Resume_tabular.xlsx")
    return FileResponse(
        path=template_path,
        filename="TEMPLATE_Resume_tabular.xlsx",
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# Calculate the ranges of experience in skills, sectors, functions, and tooling
def merge_date_ranges(ranges):
    if not ranges:
        return []

    sorted_ranges = sorted(ranges, key=lambda x: x[0])
    merged_ranges = []
    current_start, current_end = sorted_ranges[0]

    for start, end in sorted_ranges[1:]:
        if start <= current_end:
            current_end = max(current_end, end)
        else:
            merged_ranges.append((current_start, current_end))
            current_start, current_end = start, end
    merged_ranges.append((current_start, current_end))
    return merged_ranges

# Calculate the actual years of experience in skills, sectors, functions, and tooling
def calculate_experience(resume):
    experience = {"skills": {}, "sector": {}, "function": {}, "tools": {}}

    def add_experience(experience_dict, category, start_year, end_year):
        if category not in experience_dict:
            experience_dict[category] = []
        experience_dict[category].append((start_year, end_year))

    for index, job in resume.iterrows():
        start_year = int(job['Start'])
        end_year = int(job['End']) if job['End'] else datetime.now().year

        # Process skills
        for skill in job['Skills'].split(','):
            skill = skill.strip()
            add_experience(experience['skills'], skill, start_year, end_year)

        # Process sector
        sector = job['Sector'].strip()
        add_experience(experience['sector'], sector, start_year, end_year)

        # Process function
        function = job['Function'].strip()
        add_experience(experience['function'], function, start_year, end_year)

        # Process tools
        for tool in job['Tools_and_technologies_used'].split(','):
            tool = tool.strip()
            add_experience(experience['tools'], tool, start_year, end_year)

    # Merge date ranges and calculate unique years of experience
    for category in experience:
        for key in experience[category]:
            merged_ranges = merge_date_ranges(experience[category][key])
            unique_years = sum(end - start + 1 for start, end in merged_ranges)
            experience[category][key] = unique_years
    return experience

# Asynchronous function to simulate the letter creation process
async def run_letter_creation_process(resume_df, job_description, user_feedback=None, last_output=None):
    await asyncio.sleep(1)
    logger.info("Calculating experience...")
    await asyncio.sleep(1)

    # Calculate experience
    experience = calculate_experience(resume_df)

    resume_details = "\n\n".join(
        f"Company: {row['Company']}\n"
        f"Sector: {row['Sector']}\n"
        f"Function: {row['Function']}\n"
        f"Job Description Resume: {row['Job_description']}\n"
        f"Achievements: {row['Achievements']}\n"
        f"Tools and Technologies Used: {row['Tools_and_technologies_used']}\n"
        f"Skills: {row['Skills']}\n"
        f"Start: {row['Start']}\n"
        f"End: {row['End']}"
        for _, row in resume_df.iterrows()
    )

    # Prepare the experience data
    experience_input = f"""
    Skills: {experience['skills']}
    Sector: {experience['sector']}
    Function: {experience['function']}
    Tools: {experience['tools']}
    """

    # User proxy agent
    user_proxy_auto = autogen.UserProxyAgent(
        name="User_Proxy_Auto",
        human_input_mode="NEVER",
        is_termination_msg=lambda x: x.get("content", "") and x.get("content", "").rstrip().endswith("TERMINATE"),
        code_execution_config={
            "last_n_messages": 1,
            "work_dir": "tasks",
            "use_docker": False,
        },
    )

    # Matching agent
    await asyncio.sleep(1)
    logger.info("Matching resume to job description...")
    await asyncio.sleep(1)

    match_task = """
    Your task is to find the matches between the provided job description and the candidate's calculated experience. 
    Please follow these steps:

    1. Analyze the job description to extract key requirements in terms of skills, sectors, functions, and tools.
    2. Compare these requirements with the candidate's experience.
    3. Identify the most relevant matches for each category (skills, sectors, functions, tools).
    4. Provide a detailed list of these matches, including a brief explanation of how each match is relevant to the job description.
    5. Highlight any unique or standout qualifications that make the candidate a strong fit for the position.

    Reply 'MATCH_DONE' at the end of your response.
    """

    match_agent = autogen.AssistantAgent(
        name="Match_Agent",
        llm_config=llm_config,
        system_message="""
            You are an expert match maker for job applications. Your task is to meticulously analyze the provided job description 
            and the candidate's calculated experience, identifying the most relevant skills, sectors, functions, and tools that 
            match the job requirements. Your response should include a detailed list of matches along with a brief explanation 
            of how each match is relevant to the job description. Ensure to highlight any unique or standout qualifications 
            that make the candidate a strong fit for the position. 
            Reply 'MATCH_DONE' at the end of your response.
            """,
    )

    # Step 1: Find matches between job description and calculated experience
    match_task_with_experience = f"""
    Match Task: {match_task}
    Job Description: {job_description}
    Calculated Experience: {experience_input}
    """

    match_results = autogen.initiate_chats(
        [
            {
                "sender": user_proxy_auto,
                "recipient": match_agent,
                "message": match_task_with_experience,
                "summary_method": "last_msg",
                "max_turns": 1,
            }
        ]
    )

    match_result = match_results[0].summary

    if "MATCH_DONE" not in match_result:
        raise Exception("Matching process did not complete successfully.")

    # Strategy agent
    await asyncio.sleep(1)
    logger.info("Develop custom writing strategy...")
    await asyncio.sleep(1)

    strategy_task = """
    Develop a comprehensive strategy for writing exceptional motivation and cover letters, considering diverse application contexts such as job applications, university admissions, and grant proposals. The strategy should include:

    1. Methods for identifying the purpose of each letter and aligning it with the organization's objectives.
    2. Guidelines for analyzing the intended audience to tailor the content accurately.
    3. A plan for structuring these letters and personalizing them to highlight the applicant's unique skills, passions, and qualifications. Mention relevant achievements.
    4. Guidance on using persuasive language and maintaining a professional tone.
    5. Approaches for emphasizing the applicant's unique value, focusing on specific skills or experiences that distinguish them from other candidates and benefit the organization.
    6. Best practices for proofreading and editing to ensure error-free, polished documents.
    7. Very important is dynamic and captivating first opening line and compelling closing statements that make a memorable impact and urge the reader to take action.
    8. Never lie, you can enhance facts, but not make up ones. 

    Compile these guidelines into a comprehensive document.

    Reply 'STRATEGY_DONE' at the end of your response.
    """

    strategy_agent = autogen.AssistantAgent(
        name="Strategy_Agent",
        llm_config=llm_config,
        system_message="""
            You are an expert in crafting job-specific motivation letter strategies. Your role is to analyze the job description 
            and the candidate’s experience to produce a focused strategy for the letter. The strategy should include key points to highlight, 
            suggested tone and style, structuring, and other stylistic recommendations that align with the company's culture and objectives.
            Reply 'STRATEGY_DONE' at the end of your response.t that can serve as a blueprint for drafting motivation and cover letters.
            
            Reply 'STRATEGY_DONE' at the end of your response.
            """,
    )

    # Step 2: Find the best strategy for the motivation letter
    strategy_input = f"""
    Job Description: {job_description}
    Calculated Experience: {experience_input}
    Matches found: {match_result}
    Resume details: {resume_details}
    """

    strategy_results = autogen.initiate_chats(
        [
            {
                "sender": user_proxy_auto,
                "recipient": strategy_agent,
                "message": strategy_task,
                "carryover": strategy_input,
                "summary_method": "last_msg",
                "max_turns": 1,
            }
        ]
    )

    strategy_result = strategy_results[0].summary

    if "STRATEGY_DONE" not in strategy_result:
        raise Exception("Strategy generation did not complete successfully.")

    # Writer agent
    await asyncio.sleep(1)
    logger.info("Write the cover letter...")
    await asyncio.sleep(1)

    motivation_letter_task = """
    Using the provided strategy, write a compelling motivation letter based on the job description and the candidate’s experience. Ensure the letter:

    1. Highlights the most relevant skills, experiences, and achievements.
    2. Follows the outlined structure and tone as per the strategy.
    3. Includes a clear and persuasive call to action.
    4. Demonstrates the candidate's fit for the role and the company's mission and culture.
    5. Uses engaging language while maintaining a professional tone.
    6. Is tailored specifically to the job and company with relevant keywords.

    Reply 'TERMINATE' at the end of your response.
    """

    writing_agent = autogen.AssistantAgent(
        name="Writing_Agent",
        llm_config=llm_config,
        system_message="""
            You are a professional writer specializing in crafting compelling motivation letters based on a provided strategy. 
            Your task is to draft the motivation letter by following the structured outline, tone, and emphasis detailed in the strategy. 
            Focus on highlighting the candidate’s relevant achievements, skills, and experiences to make them a strong fit for the job.
            Reply 'TERMINATE' at the end of your response.

            Reply 'TERMINATE' at the end of your response.
            """,
    )

    # Step 3: Write a motivation letter
    motivation_letter_input = f"""
    Matches found: {match_result}
    Strategy suggested: {strategy_result}
    Resume Details: {resume_details}
    """

    if last_output:
        motivation_letter_input += f"\n\nPrevious Output: {last_output}"
    if user_feedback:
        motivation_letter_input += f"\n\nUser Feedback: {user_feedback}"

    motivation_letter_results = autogen.initiate_chats(
        [
            {
                "sender": user_proxy_auto,
                "recipient": writing_agent,
                "message": motivation_letter_task,
                "carryover": motivation_letter_input,
                "max_turns": 1,
            }
        ]
    )

    motivation_letter_result = motivation_letter_results[0].summary

    # Quality check agent
    await asyncio.sleep(1)   
    logger.info("Check quality of the letter to potentially rewrite...")
    await asyncio.sleep(1)
    
    quality_check_task = """
    Review the motivation letter for the following criteria:

    1. Clarity: Ensure the content is easy to understand and free from ambiguity.
    2. Coherence: Check that the ideas flow logically and the letter is well-structured.
    3. Grammar: Correct any grammatical, spelling, or punctuation errors.
    4. Adherence to Guidelines: Ensure the letter follows the provided structure and strategy guidelines.
    5. Formatting: Verify the document is well-formatted, with proper sections and bullet points where necessary.
    6. Overall Quality: Assess the letter's overall quality, including the strength of the opening, alignment with the company’s mission and goals, and the effectiveness of the call to action.
    7. Highlighting Achievements: Ensure key achievements and quantifiable results are clearly highlighted.
    8. Tailoring: Verify that the content is specifically tailored to the job and company, demonstrating an understanding of their needs and culture.
    9. Keywords: Check that relevant keywords from the job description are included to ensure it passes through Applicant Tracking Systems (ATS).
    10. Enthusiasm and Cultural Fit: Ensure the candidate's enthusiasm for the role and cultural fit with the company are effectively showcased.

    Provide detailed feedback and necessary corrections. Reply 'QUALITY_DONE' at the end of your review. If revisions are needed, provide specific feedback and reply 'REWRITE_NEEDED' at the very end.
    """

    quality_check_agent = autogen.AssistantAgent(
        name="Quality_Check_Agent",
        llm_config=llm_config,
        system_message="""
            You are a quality assurance specialist for professional documents. 
            Your task is to review the motivation letter for clarity, coherence, grammar, and adherence to the provided structure and guidelines. 
            Ensure the letter is well-formatted, with proper sections and bullet points where necessary. 
            Provide detailed feedback and necessary corrections to enhance the letter's overall quality and impact.

            Reply 'QUALITY_DONE' at the end of your review. If revisions are needed, provide specific feedback and reply 'REWRITE_NEEDED' at the very end.
            """,
    )

    # Step 4: Quality check the motivation letter
    quality_check_input = f"""
    Motivation Letter: {motivation_letter_result}
    """

    quality_check_results = autogen.initiate_chats(
        [
            {
                "sender": user_proxy_auto,
                "recipient": quality_check_agent,
                "message": quality_check_task,
                "carryover": quality_check_input,
                "summary_method": "last_msg",
                "max_turns": 1,
            }
        ]
    )

    quality_check_result = quality_check_results[0].summary

    # Ensure the quality check is complete
    if "QUALITY_DONE" not in quality_check_result and "REWRITE_NEEDED" not in quality_check_result:
        raise Exception("Quality check did not complete successfully.")

    # If quality check recommends rewriting, additional feedback is incorporated
    if "REWRITE_NEEDED" in quality_check_result:
        feedback = quality_check_result.replace('REWRITE_NEEDED', '').strip()
        if user_feedback:
            feedback += f"\nUser Feedback: {user_feedback}"
        motivation_letter_results = autogen.initiate_chats(
            [
                {
                    "sender": user_proxy_auto,
                    "recipient": writing_agent,
                    "message": motivation_letter_task,
                    "carryover": motivation_letter_result + f"\n\nQuality Check Feedback: {feedback}",
                    "summary_method": "last_msg",
                    "max_turns": 1,
                }
            ]
        )
        motivation_letter_result = motivation_letter_results[0].summary

    return motivation_letter_result

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
