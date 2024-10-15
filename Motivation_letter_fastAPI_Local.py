from fastapi import FastAPI, File, UploadFile, Form, Depends
from fastapi.responses import RedirectResponse
from fastapi.responses import StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import Optional
from io import BytesIO
from docx import Document
import pandas as pd
import os
import datetime
import autogen
from fastapi.middleware.cors import CORSMiddleware
from transformers import AutoModelForCausalLM, AutoTokenizer, GPT2Tokenizer, GPTNeoForCausalLM, GPTNeoConfig
import torch


app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Or replace "*" with your Netlify URL or other trusted domains
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve static files (HTML, CSS, etc.)
app.mount("/static", StaticFiles(directory="static"), name="static")


# # Configuration setup
# config_list = autogen.config_list_from_json(env_or_file="OAI_CONFIG_LIST")
# filtered_config_list = [item for item in config_list if item["model"] == "gpt-4"]
# # print(filtered_config_list)
# llm_config = {"config_list": filtered_config_list}


# Load the GPT-J model and tokenizer once at startup
# model_name = "EleutherAI/gpt-j-6B"
# model_name = "EleutherAI/gpt-neo-2.7B"
model_name = "EleutherAI/gpt-neo-1.3B"
tokenizer = AutoTokenizer.from_pretrained(model_name)
# model = AutoModelForCausalLM.from_pretrained(model_name, device_map="auto", torch_dtype=torch.float16).to("cuda")
# model = AutoModelForCausalLM.from_pretrained(model_name, device_map="auto", torch_dtype=torch.float16).to("cpu")


# # Load the tokenizer
# tokenizer = GPT2Tokenizer.from_pretrained(model_name)

# # Load the GPT-Neo model with 8-bit quantization for memory efficiency
# model = GPTNeoForCausalLM.from_pretrained(
#     model_name,
#     load_in_8bit=True,  # Use 8-bit quantization
#     device_map="auto"  # Automatically allocate layers to available devices
# )

# Modify the model's configuration for longer sequence lengths
config = GPTNeoConfig.from_pretrained(model_name)
config.max_position_embeddings = 4096  # Increase max token limit (from 2048 to 4096)

# Load the model with the modified configuration and ignore mismatched sizes
model = GPTNeoForCausalLM.from_pretrained(
    "EleutherAI/gpt-neo-1.3B", 
    config=config, 
    ignore_mismatched_sizes=True  # Ignore size mismatch for position embeddings
)

# Now manually resize the position embedding layer
# Get the hidden size from the model configuration
hidden_size = config.hidden_size

# Resize the position embedding layer (wpe) to match the new max_position_embeddings
old_wpe = model.transformer.wpe  # Old position embeddings
new_wpe = torch.nn.Embedding(config.max_position_embeddings, hidden_size)  # New position embeddings

# Copy the weights from the old position embeddings to the new embeddings
new_wpe.weight.data[:old_wpe.weight.size(0)] = old_wpe.weight.data

# Replace the old position embeddings with the new resized ones
model.transformer.wpe = new_wpe

# Move the model to CPU
model = model.to("cpu")


def generate_response(prompt: str, max_tokens: int = 4000):
    # Tokenize the input prompt
    inputs = tokenizer(prompt, return_tensors="pt").to("cpu")

    # Generate attention mask
    attention_mask = inputs['attention_mask']
    
    # Generate response using the GPT-J model
    outputs = model.generate(
        inputs["input_ids"], 
        max_length=max_tokens, 
        pad_token_id=tokenizer.eos_token_id, 
        attention_mask=attention_mask
    )
    
    # Decode the output tokens
    response = tokenizer.decode(outputs[0], skip_special_tokens=True)
    return response

@app.get("/generate/")
def generate(prompt: str):
    response = generate_response(prompt)
    return {"response": response}



# Root route that redirects to the static index.html
@app.get("/")
async def serve_index():
    return RedirectResponse(url="/static/index.html")

# Endpoint for resume upload and job description input
@app.post("/generate-letter/")
async def generate_motivation_letter(
    resume_file: UploadFile = File(...), job_description: str = Form(...)
):
    # Read resume file and cache it in-memory for the user session
    resume_df = pd.read_excel(resume_file.file, sheet_name='Jobs')
    
    # # Store resume and job description in user cache
    # user_cache['resume'] = resume_df
    # user_cache['job_description'] = job_description

    # Call the multi-agent system to generate the letter
    # motivation_letter_result = run_letter_creation_process(user_cache['resume'], user_cache['job_description'])
    motivation_letter_result = run_letter_creation_process(resume_df, job_description)

    # Generate motivation letter as a Word document in-memory
    document = Document()
    document.add_heading('Motivation Letter', level=1)
    document.add_paragraph(motivation_letter_result)
    
    # Prepare the file for download without storing it locally
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    # Provide file download response using StreamingResponse
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=motivation_letter.docx"}
    )


# Calculate the ranges of experiene in skills, sectors, functions and tooling. E.g. [(2015, 2017), (2013, 2014)]
def merge_date_ranges(ranges):
    if not ranges:
        return []

    sorted_ranges = sorted(ranges, key=lambda x: x[0])
    merged_ranges = []
    current_start, current_end = sorted_ranges[0]

    for start, end in sorted_ranges[1:]:
        if start <= current_end:
            current_end = max(current_end, end)
        else:
            merged_ranges.append((current_start, current_end))
            current_start, current_end = start, end
    merged_ranges.append((current_start, current_end))
    
    # print(f"Merged ranges: {merged_ranges}")  # Debug statement
    return merged_ranges

# Calculate the actual years of experience in skills, sectors, functions and tooling
def calculate_experience(resume):
    experience = {
        "skills": {},
        "sector": {},
        "function": {},
        "tools": {}
    }

    def add_experience(experience_dict, category, start_year, end_year):
        if category not in experience_dict:
            experience_dict[category] = []
        experience_dict[category].append((start_year, end_year))

    for index, job in resume.iterrows():
        start_year = int(job['Start'])
        end_year = int(job['End']) if job['End'] else datetime.now().year

        # Process skills
        for skill in job['Skills'].split(','):
            skill = skill.strip()
            add_experience(experience['skills'], skill, start_year, end_year)

        # Process sector
        sector = job['Sector'].strip()
        add_experience(experience['sector'], sector, start_year, end_year)

        # Process function
        function = job['Function'].strip()
        add_experience(experience['function'], function, start_year, end_year)

        # Process tools
        for tool in job['Tools_and_technologies_used'].split(','):
            tool = tool.strip()
            add_experience(experience['tools'], tool, start_year, end_year)

    # Merge date ranges and calculate unique years of experience
    for category in experience:
        for key in experience[category]:
            # print(f"Category: {category}, Key: {key}, Ranges before merge: {experience[category][key]}")  # Debug statement
            merged_ranges = merge_date_ranges(experience[category][key])
            unique_years = sum(end - start + 1 for start, end in merged_ranges)
            experience[category][key] = unique_years
            # print(f"Category: {category}, Key: {key}, Unique years: {unique_years}")  # Debug statement

    return experience

def run_letter_creation_process(resume_df, job_description):
    # Calculate experience
    experience = calculate_experience(resume_df)

    # Print the calculated experience
    # print("Experience calculated from Excel file:")
    # print(experience)

    resume_details = "\n\n".join(
        f"Company: {row['Company']}\n"
        f"Sector: {row['Sector']}\n"
        f"Function: {row['Function']}\n"
        f"Job Description Resume: {row['Job_description']}\n"
        f"Achievements: {row['Achievements']}\n"
        f"Tools and Technologies Used: {row['Tools_and_technologies_used']}\n"
        f"Skills: {row['Skills']}\n"
        f"Start: {row['Start']}\n"
        f"End: {row['End']}"
        for _, row in resume_df.iterrows()
    )

    # Prepare the experience data
    experience_input = f"""
    Skills: {experience['skills']}
    Sector: {experience['sector']}
    Function: {experience['function']}
    Tools: {experience['tools']}
    """

    # Define tasks for agents
    # job_description = input("Please enter the job description: ") # This is the start of the 'conversation' where the user pastes in a job-description
    match_task = """
    Your task is to find the matches between the provided job description and the candidate's calculated experience. 
    Please follow these steps:

    1. Analyze the job description to extract key requirements in terms of skills, sectors, functions, and tools.
    2. Compare these requirements with the candidate's experience.
    3. Identify the most relevant matches for each category (skills, sectors, functions, tools).
    4. Provide a detailed list of these matches, including a brief explanation of how each match is relevant to the job description.
    5. Highlight any unique or standout qualifications that make the candidate a strong fit for the position.

    Reply 'MATCH_DONE' at the end of your response.
    """

    match_maker = """
            You are an expert match maker for job applications. Your task is to meticulously analyze the provided job description 
            and the candidate's calculated experience, identifying the most relevant skills, sectors, functions, and tools that 
            match the job requirements. Your response should include a detailed list of matches along with a brief explanation 
            of how each match is relevant to the job description. Ensure to highlight any unique or standout qualifications 
            that make the candidate a strong fit for the position. 
            Reply 'MATCH_DONE' at the end of your response.
    """
    

    match_task_with_experience = f"""
    Match Agent: {match_maker}
    Match Task: {match_task}
    Job Description: {job_description}
    Calculated Experience: {experience_input}
    """

    # Use the local GPT-J model to generate the response
    match_result = generate_response(match_task_with_experience)

    return match_result


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
